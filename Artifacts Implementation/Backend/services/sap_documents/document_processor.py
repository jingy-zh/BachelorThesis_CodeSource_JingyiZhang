"""
Document processing module for extracting and manipulating content from DOCX files.

This module provides functionality for processing Word documents, including
extracting tables, images, and text content, splitting text into chunks,
removing tables and images, and replacing images with S3 links. It handles
various document elements while maintaining document structure information.
"""
import base64
import io

import pandas as pd
from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from langchain.text_splitter import RecursiveCharacterTextSplitter
from loguru import logger
from PIL import Image

from .config import sap_documents_config


class DocumentProcessor:
	"""
	Processes DOCX documents to extract and manipulate their content.

	This class provides methods to extract tables, images, and text from
	Word documents, resize images, split text into manageable chunks,
	remove tables and images from documents, and replace images with S3 links.
	"""
	def __init__(self):
		self.text_splitter = RecursiveCharacterTextSplitter(
			chunk_size=sap_documents_config.CHUNK_SIZE,
			chunk_overlap=sap_documents_config.CHUNK_OVERLAP,
		)
		self.max_image_size = sap_documents_config.MAX_IMAGE_SIZE

	def extract_tables(self, doc, file_name, is_save):
		"""
		Extract tables from a Word document and convert them to markdown format.

		Processes each table in the document, converts it to a pandas DataFrame,
		optionally saves it as a CSV file, and returns all tables in markdown format.

		Args:
		    doc: Word document containing tables to extract
		    file_name: Base name for saved table files
		    is_save: Whether to save extracted tables as CSV files

		Returns:
		    list: List of tables in markdown format
		"""
		tables = []
		for i, table in enumerate(doc.tables):
			data = []
			for row in table.rows:
				row_data = [cell.text for cell in row.cells]
				data.append(row_data)

			df = pd.DataFrame(data[1:], columns=data[0])

			# Save table to CSV file
			if is_save is True:
				csv_filename = f'{sap_documents_config.TABLE_FOLDER}/{file_name}_table_{i+1}.csv'
				df.to_csv(csv_filename, index=False)

			markdown_table = df.to_markdown()
			tables.append(str(markdown_table))

		return tables

	def extract_images(self, doc, file_name, is_save):
		"""
		Extract, resize, and encode images from a Word document.

		Processes all images in the document, resizes them to the configured maximum
		size while maintaining aspect ratio, skips images smaller than the minimum
		dimensions, and optionally saves them to disk.

		Args:
		    doc: Word document containing images to extract
		    file_name: Base name for saved image files
		    is_save: Whether to save extracted images to disk

		Returns:
		    list: List of base64-encoded image strings
		"""
		images = []
		for rel in doc.part.rels.values():
			if 'image' in rel.reltype:
				try:
					image = rel.target_part.blob
					image_file = io.BytesIO(image)

					# Open the image using Pillow
					with Image.open(image_file) as img:
						# Skip image if it is too small
						if img.width < sap_documents_config.MIN_WIDTH and img.height < sap_documents_config.MIN_HEIGHT:
							logger.info(f'Skipping image in {file_name}: {img.width}x{img.height}')
							continue

						# Resize the image while maintaining aspect ratio
						img.thumbnail(self.max_image_size)

						# Save the resized image to a bytes buffer
						output_buffer = io.BytesIO()
						img.save(output_buffer, format=img.format)
						resized_image = output_buffer.getvalue()

					encoded_image = base64.b64encode(resized_image).decode('utf-8')
					images.append(encoded_image)

					if is_save is True:
						image_filename = f'{sap_documents_config.IMAGE_FOLDER}/{file_name}_{len(images)}.png'
						with open(image_filename, 'wb') as f:
							f.write(resized_image)
				except Exception as e:
					# This catches the error for external images
					logger.error(f'Skipping external image in {file_name}: {e}')
					continue
		return images

	def remove_tables_and_images(self, doc):
		"""
		Remove tables and images from a Word document, keeping only the text content.

		This method iterates through the document's content, identifying paragraphs and tables.
		It creates a new document containing only the text from paragraphs, effectively
		removing all tables and images.

		Args:
			doc (Document): The input Word document.

		Returns:
			Document: A new Document object containing only the text content.

		Note:
			This method uses a nested function `iter_block_items` to iterate through
			the document's content, handling both the main document body and table cells.
		"""

		def iter_block_items(parent):
			if isinstance(parent, _Document):
				parent_elm = parent.element.body
			elif isinstance(parent, _Cell):
				parent_elm = parent._tc
			else:
				raise ValueError("Something's not right")

			for child in parent_elm.iterchildren():
				if isinstance(child, CT_P):
					yield Paragraph(child, parent)
				elif isinstance(child, CT_Tbl):
					yield Table(child, parent)

		new_doc = Document()
		for block in iter_block_items(doc):
			if isinstance(block, Paragraph):
				new_doc.add_paragraph(block.text)

		return new_doc

	def extract_texts(self, doc):
		"""
		Extract and split text content from a Word document.

		Extracts text from all paragraphs, joins them, and splits the text into
		chunks of the configured size with the configured overlap.

		Args:
		    doc: Word document to extract text from

		Returns:
		    list: List of text chunks, each approximately the configured size
		"""
		contents = []
		for paragraph in doc.paragraphs:
			contents.append(paragraph.text)

		joined_texts = ' '.join(contents)
		texts_4k_token = self.text_splitter.split_text(joined_texts)

		return texts_4k_token

	def extract_content_with_positions_and_styles(self, doc):
		"""
		Extract document content preserving positions, styles, and embedded objects.

		Processes paragraphs and runs in the document, capturing text content,
		embedded images, and their positions in the document structure. Skips
		images smaller than the configured minimum dimensions.

		Args:
		    doc: Document object to extract content from

		Returns:
		    list: Hierarchical structure of document content with positions and styles
		"""
		content = []

		for i, paragraph in enumerate(doc.paragraphs):
			runs = []
			for run in paragraph.runs:
				runs.append(('text', run, run.text))
				if 'graphic' in run._element.xml:
					for drawing in run._element.xpath('.//a:blip'):
						rId = drawing.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
						image_part = doc.part.related_parts[rId]
						image_stream = io.BytesIO(image_part.blob)
						image = Image.open(image_stream)
						if image.width < sap_documents_config.MIN_WIDTH and image.height < sap_documents_config.MIN_HEIGHT:
							logger.info(f'Skipping image: {image.width}x{image.height}')
							continue

						runs.append(('image', run, image))
			content.append(('paragraph', i, runs))

		return content

	def replace_images_with_s3_links(self, content, file_path, s3_image_uploader):
		"""
		Replace images in document content with S3 links.

		Processes each image in the document, resizes it, uploads it to S3,
		and replaces the image in the document with a link to the S3 location.

		Args:
		    content: Document content with position and style information
		    file_path: Path to the document file
		    s3_image_uploader: Service for uploading images to S3

		Returns:
		    dict: Dictionary containing encoded images and their file names
		"""
		encoded_images = []
		file_name_images = []

		doc_file_name = file_path.split('/')[-1].split('.')[0]
		for _, pos, runs in content:
			for index, (run_type, run, img) in enumerate(runs):
				if run_type == 'image':
					image_file_name = f'{doc_file_name}_{pos}_{index}'
					try:
						# Resize the image while maintaining aspect ratio
						img.thumbnail(self.max_image_size)

						# Save the resized image to a bytes buffer
						output_buffer = io.BytesIO()
						img.save(output_buffer, format='PNG')
						output_buffer.seek(0)
						encoded_image = base64.b64encode(output_buffer.read()).decode('utf-8')

						# Upload the image to S3 and get the URL
						filename = f'{image_file_name}.png'
						s3_image_uploader.save_image('image/png', encoded_image, filename)
						logger.info(f'Saved image to S3: {filename}')
						combined = f'[(VSDA_S3_KEY:{filename})]'
						run.text = combined

						# Append the encoded image and the file name image to the lists
						encoded_images.append(encoded_image)
						file_name_images.append(filename)
					except Exception as e:
						logger.error(f'Error replacing image with S3 link: {e}')

		extracted_images = {'encoded_images': encoded_images, 'file_name_images': file_name_images}
		return extracted_images

	def process_document(self, doc, output_file, file_name, is_save, s3_image_uploader):
		"""
		Process a document to extract tables, images, and text.

		Comprehensive document processing that extracts tables, images, and text,
		replacing images with S3 links, removing tables and images from the text,
		and optionally saving the modified document.

		Args:
		    doc: Word document to process
		    output_file: Path where the modified document should be saved
		    file_name: Base name for any saved files
		    is_save: Whether to save the modified document and extracted content
		    s3_image_uploader: Service for uploading images to S3

		Returns:
		    tuple: Contains extracted tables, images, and text chunks
		"""
		extracted_tables = self.extract_tables(doc, file_name, is_save)
		# extracted_images = self.extract_images(doc, file_name, is_save)
		content = self.extract_content_with_positions_and_styles(doc=doc)
		extracted_images = self.replace_images_with_s3_links(content, file_name, s3_image_uploader)
		new_doc = self.remove_tables_and_images(doc)
		extracted_text = self.extract_texts(doc=new_doc)

		logger.info(f'Extracted {len(extracted_text)} chunk texts')
		logger.info(f'Extracted {len(extracted_tables)} tables')
		logger.info(f'Extracted {len(extracted_images["encoded_images"])} images')

		if is_save is True:
			new_doc.save(output_file)
			logger.info(f'Saved modified document as {output_file}')

		return extracted_tables, extracted_images, extracted_text
